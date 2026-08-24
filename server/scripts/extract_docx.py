import docx

doc = docx.Document(r'C:\Users\LA TCHAUX HOTEL\Downloads\DSF_OHADA_Regles_Metier.docx')

with open('extracted_dsf_regles_metier.txt', 'w', encoding='utf-8') as f:
    f.write("=== REGLES METIER DSF OHADA ===\n\n")
    for p in doc.paragraphs:
        if p.text.strip():
            f.write(p.text + "\n")
    
    f.write("\n=== TABLES ===\n")
    for i, table in enumerate(doc.tables):
        f.write(f"\n--- TABLE {i+1} ---\n")
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            f.write(" | ".join(cells) + "\n")

print("Extracted successfully into extracted_dsf_regles_metier.txt")
